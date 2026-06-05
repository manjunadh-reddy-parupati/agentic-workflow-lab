"""
================================================================================
 Document Governance Workflow  —  Standalone Python Edition
================================================================================

A single-file, fully self-contained agentic workflow built on the Microsoft
Agent Framework (Python) that checks a document against approved reference
content.

WHAT IT DOES
------------
1. Reads a .docx (or .md / .txt fallback) document from the ./input folder.
2. Splits it into sections by heading style (python-docx).
3. For each section, runs a graph workflow:
       DocumentSegmenter -> SegmentDispatcher -> ReferenceRetriever -> AlignmentReviewer -> (loop) -> ReportCompiler
   - ReferenceRetriever : Local Verba (Weaviate) vector retrieval of the approved "source of truth".
   - AlignmentReviewer  : Azure OpenAI (multimodal) checks each section for DIRECT
                          contradictions against the retrieved reference chunks.
   - ReportCompiler     : prioritises discrepancies and builds the final QC report.
4. Writes the QC report as JSON to the ./output folder.

HOW TO RUN
----------
    python doc_governance_workflow.py

The Main Document is parsed locally (python-docx) and NEVER vectorized; each of
its sections is sent as a live query to Verba, which retrieves matching chunks
from the pre-indexed Source-of-Truth document.

DEPENDENCIES (already expected in the environment)
--------------------------------------------------
    agent-framework            (Microsoft Agent Framework - Python)
    python-docx                (DOCX parsing)
    Pillow                     (image resize for vision model)
    httpx                      (HTTP client for the local Verba RAG service)
    openai / azure             (LLM via agent framework azure client)
================================================================================
"""

import asyncio
import io
import json
import os
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from typing_extensions import Never

# --- Centralized configuration -------------------------------------------------
# Load every setting from the single project config file (configs/.env) so that
# nothing sensitive (API keys, endpoints) is hard-coded in source. Works whether
# this module is imported by the backend or run directly as a script.
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parent.parent / "configs" / ".env")

# Ensure UTF-8 console output on Windows (emoji/box-drawing chars in logs).
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")  # type: ignore[attr-defined]
    except (AttributeError, ValueError):
        pass

# --- Microsoft Agent Framework -------------------------------------------------
from agent_framework import (
    ChatMessage,
    Content,
    Executor,
    WorkflowBuilder,
    WorkflowContext,
    WorkflowEvent,
    WorkflowOutputEvent,
    handler,
)
from agent_framework.azure import AzureOpenAIChatClient

# --- Local RAG (Verba over Weaviate) -------------------------------------------
# Supports both module import (backend: workflows.doc_governance_workflow)
# and direct script execution (python doc_governance_workflow.py).
try:
    from .verba_client import VerbaClient
except ImportError:  # running as a standalone script
    from verba_client import VerbaClient

# --- DOCX + image processing ---------------------------------------------------
from PIL import Image
from docx import Document as DocxDocument
from docx.document import Document as _DocxDocumentClass
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


# ==============================================================================
# region CONFIGURATION  (all values hard-coded here)
# ==============================================================================
class CONFIG:
    # === Folders (input/output live at the project root, shared by the whole app) ===
    SCRIPT_DIR = Path(__file__).resolve().parent
    PROJECT_ROOT = SCRIPT_DIR.parent
    INPUT_FOLDER = PROJECT_ROOT / "input"
    OUTPUT_FOLDER = PROJECT_ROOT / "output"
    DIAGRAM_OUTPUT_FOLDER = PROJECT_ROOT / "diagram-outputs"

    # === Document fallback (used only if no .docx is found in input/) ===
    MARKDOWN_FALLBACK_PATH = INPUT_FOLDER / "MD" / "csr.txt"

    # === Section processing ===
    MAX_SECTION_CHARACTERS = 40000      # split sections larger than this
    MAX_SEARCH_QUERY_CHARACTERS = 8000  # max characters per Verba query (embedding token buffer)
    VERBA_TOP_K = int(os.environ.get("VERBA_TOP_K", "3"))  # reference chunks retrieved per query

    # === Image processing ===
    ENABLE_IMAGE_ANALYSIS = True
    MAX_IMAGE_PIXELS = 2048             # resize images larger than this (long edge)

    # === Local RAG (Weaviate) — connection details come from env (see VerbaClient.from_env) ===
    # WEAVIATE_HOST       (default localhost)
    # WEAVIATE_HTTP_PORT  (default 8080)
    # WEAVIATE_GRPC_PORT  (default 50051)
    # RAG_COLLECTION      (default SourceOfTruth)

    # === Azure OpenAI (loaded from configs/.env — never hard-code secrets) ===
    AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT", "")
    AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY", "")
    AZURE_OPENAI_DEPLOYMENT = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-5.2")
    AZURE_OPENAI_MODEL = os.environ.get("AZURE_OPENAI_MODEL", "gpt-5.2")
    AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2024-12-01-preview")

    # === Review Agent ===
    REVIEW_AGENT_NAME = "DocGovernance_AlignmentReviewer"
    REVIEW_AGENT_INSTRUCTIONS = (
        "You are an expert Quality Control (QC) Agent that checks documents against approved reference content.\n\n"
        "IMPORTANT - SOURCE OF TRUTH:\n"
        "- The APPROVED REFERENCE CONTENT is the AUTHORITATIVE SOURCE OF TRUTH. It is pre-approved, verified content.\n"
        "- The SECTION CONTENT is the DOCUMENT BEING REVIEWED.\n"
        "- Your job is to find ONLY DIRECT CONTRADICTIONS where the document says something OPPOSITE to what the reference content says.\n\n"
        "CRITICAL RULES - WHAT TO FLAG vs WHAT TO IGNORE:\n\n"
        "\u2705 FLAG (Discrepancy) - ONLY these cases:\n"
        "- Document says 'The service is free' but the reference says 'The service costs $20' \u2192 CONTRADICTION, flag it\n"
        "- Document says 'Accuracy was 85%' but the reference says 'Accuracy was 65%' \u2192 CONTRADICTION, flag it  \n"
        "- Document says 'No errors occurred' but the reference says 'Errors occurred' \u2192 CONTRADICTION, flag it\n"
        "- Document says 'Throughput increased' but the reference says 'Throughput decreased' \u2192 CONTRADICTION, flag it\n"
        "- IMAGE shows 'wireless device' but the reference says 'wired connection required' \u2192 CONTRADICTION, flag it\n\n"
        "\u274c DO NOT FLAG - These are NOT discrepancies:\n"
        "- Document contains information NOT mentioned in the reference content \u2192 This is NORMAL (document is larger than the reference)\n"
        "- Document section covers topics not in the reference content \u2192 ALIGNED (silence is not a conflict)\n"
        "- Document section has operational details (dates, locations, names) not in the reference content \u2192 ALIGNED (extra details are fine)\n"
        "- The reference content mentions something that this document section doesn't cover \u2192 NOT a discrepancy (section may cover different topic)\n"
        "- Document provides more detail than the reference content on the same topic \u2192 ALIGNED (more detail is fine, only flag if it CONTRADICTS)\n\n"
        "THE CORE LOGIC:\n"
        "1. Read the document section content AND analyze any images provided\n"
        "2. Check if ANY statement in the section OR any visual content in images DIRECTLY CONTRADICTS the reference content\n"
        "3. A contradiction means: SAME TOPIC + OPPOSITE/CONFLICTING claims\n"
        "4. If no contradictions exist \u2192 isAligned = true\n"
        "5. Silence or absence of information is NEVER a discrepancy\n\n"
        "CRITICAL: Output ONLY valid JSON matching this EXACT structure (no markdown, no prefixes, no explanation):\n"
        "{\n"
        '    "isAligned": true/false,\n'
        '    "confidence": 0.0 to 1.0,\n'
        '    "severity": "Critical" | "Medium" | "Low",\n'
        '    "discrepancies": [\n'
        "        {\n"
        '            "statementInDocument": "EXACT VERBATIM QUOTE from text or visible text in image",\n'
        '            "conflictingReference": "EXACT VERBATIM QUOTE from the APPROVED REFERENCE CONTENT",\n'
        '            "issueDescription": "Explain the DIRECT CONTRADICTION",\n'
        '            "isFromImage": false,\n'
        '            "imageDescription": "",\n'
        '            "imageReference": ""\n'
        "        }\n"
        "    ],\n"
        '    "reasoning": "Brief explanation"\n'
        "}\n\n"
        "\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\n"
        "\U0001f4f7 IMAGE DISCREPANCY FORMAT - When contradiction is found IN AN IMAGE:\n"
        "\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\n\n"
        "For image-based contradictions, use THIS format:\n"
        "{\n"
        '    "statementInDocument": "The exact text visible in the image (e.g., \'battery-powered wireless unit\')",\n'
        '    "conflictingReference": "EXACT quote from the APPROVED REFERENCE CONTENT it contradicts",\n'
        '    "issueDescription": "Explain how the image content contradicts the reference content",\n'
        '    "isFromImage": true,\n'
        '    "imageDescription": "Describe what you see: e.g., \'Image shows a product box and device with text stating battery operation\'",\n'
        '    "imageReference": "Image 1 on Page X"\n'
        "}\n\n"
        "EXAMPLE IMAGE DISCREPANCY:\n"
        "If image shows: Product box labeled 'Model X' with text 'battery powered'\n"
        "And the reference content says: 'Model X is mains-powered'\n"
        "Then:\n"
        "{\n"
        '    "statementInDocument": "battery powered",\n'
        '    "conflictingReference": "Model X is mains-powered",\n'
        '    "issueDescription": "Image shows the product as battery powered while the reference content states it is mains-powered",\n'
        '    "isFromImage": true,\n'
        '    "imageDescription": "Product image showing a box labeled Model X with text describing battery operation",\n'
        '    "imageReference": "Image 1 on Page 1"\n'
        "}\n\n"
        "\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\n"
        "CRITICAL FIELD DEFINITIONS FOR TEXT DISCREPANCIES:\n"
        "\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\n\n"
        "\U0001f4c4 \"statementInDocument\" = The problematic text FROM THE DOCUMENT SECTION BEING REVIEWED\n"
        "   - For TEXT: VERBATIM copy-paste from the 'SECTION TO VALIDATE' content\n"
        "   - For IMAGES: The exact text visible in the image that contradicts\n"
        "   - This text will be used to SEARCH and HIGHLIGHT in the original Word document\n"
        "   - NEVER put reference content here\n\n"
        "\U0001f4cb \"conflictingReference\" = The approved text FROM THE REFERENCE CONTENT that the document contradicts\n"
        "   - This MUST be a VERBATIM copy-paste from the 'APPROVED REFERENCE CONTENT' provided below\n"
        "   - This shows what the CORRECT/APPROVED content should be\n\n"
        "\u26a0\ufe0f COMMON MISTAKE TO AVOID:\n"
        "   WRONG: statementInDocument contains reference text, conflictingReference contains document text\n"
        "   RIGHT: statementInDocument contains DOCUMENT/IMAGE text, conflictingReference contains REFERENCE text\n\n"
        "VALIDATION RULES:\n"
        "- For text discrepancies: statementInDocument MUST appear in the SECTION TO VALIDATE text\n"
        "- For image discrepancies: statementInDocument should be the visible text from the image, set isFromImage=true\n"
        "- conflictingReference MUST appear in the APPROVED REFERENCE CONTENT text\n"
        "\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\n\n"
        "Severity Guidelines (ONLY for actual contradictions):\n"
        "- Critical: Document/image directly contradicts the reference content on core facts, numbers, dates, specifications, or primary claims\n"
        "- Medium: Document contradicts the reference content on supporting details, secondary information, or contextual claims\n"
        "- Low: Minor wording differences, phrasing variations, or formatting inconsistencies that don't change meaning\n\n"
        "DEFAULT TO ALIGNED: If you're unsure whether something is a contradiction, it probably isn't. \n"
        "Most sections will be ALIGNED because they either match the reference content OR discuss topics not covered by the reference content.\n\n"
        "Remember: Only flag DIRECT CONTRADICTIONS. Silence, omission, or extra details are NOT discrepancies."
    )


_CFG = CONFIG()
# endregion


# ==============================================================================
# region DATA MODELS
# ==============================================================================
@dataclass
class ParagraphInfo:
    page_no: int = 0
    para_index: int = 0
    para_text: str = ""
    character_count: int = 0


@dataclass
class ImageInfo:
    data: bytes = b""
    mime_type: str = "image/png"
    page_no: int = 0
    image_index: int = 0
    alt_text: str = ""
    original_width: int = 0
    original_height: int = 0
    was_resized: bool = False


@dataclass
class ParsedSection:
    section_name: str = ""
    paragraphs: list[ParagraphInfo] = field(default_factory=list)
    images: list[ImageInfo] = field(default_factory=list)


@dataclass
class ProcessableSection:
    id: str = ""
    title: str = ""
    content: str = ""
    character_count: int = 0
    is_split: bool = False
    split_index: int = 0
    total_splits: int = 1
    paragraphs: list[ParagraphInfo] = field(default_factory=list)
    start_page: int = 0
    end_page: int = 0
    images: list[ImageInfo] = field(default_factory=list)


@dataclass
class ExtractedSections:
    sections: list[ProcessableSection] = field(default_factory=list)
    total_count: int = 0
    user_friendly_message: str = ""


@dataclass
class ReferenceSnippet:
    id: str = ""
    content: str = ""
    source: str = ""


@dataclass
class Discrepancy:
    statement_in_document: str = ""
    conflicting_reference: str = ""
    issue_description: str = ""
    is_from_image: bool = False
    image_description: str = ""
    image_reference: str = ""


@dataclass
class ValidationResult:
    section_id: str = ""
    section_title: str = ""
    is_aligned: bool = True
    confidence: float = 0.0
    severity: str = "Medium"  # Critical | Medium | Low
    discrepancies: list[Discrepancy] = field(default_factory=list)
    reasoning: str = ""


# === Workflow message envelopes ===
@dataclass
class DocumentJob:
    file_path: Optional[str] = None
    selection_text: Optional[str] = None
    is_full_document: bool = True
    user_selection: bool = False
    selected_sections: Optional[list[str]] = None
    resume_from_section: int = 0

    @property
    def is_valid(self) -> bool:
        return bool((self.file_path or "").strip()) or bool((self.selection_text or "").strip())


@dataclass
class DispatchSignal:
    initial_sections: Optional[ExtractedSections] = None
    loop_back_signal: Optional[ValidationResult] = None
    resume_from_section: int = 0


@dataclass
class DispatchResult:
    current_section: Optional[ProcessableSection] = None
    is_complete: bool = False
    processed_count: int = 0
    total_count: int = 0
    all_validation_results: list[ValidationResult] = field(default_factory=list)
    was_paused: bool = False


@dataclass
class RetrievalResult:
    section: ProcessableSection = field(default_factory=ProcessableSection)
    retrieved_references: list[ReferenceSnippet] = field(default_factory=list)


# Severity ordering — mirrors the C# enum (Critical=0, Medium=1, Low=2).
SEVERITY_RANK = {"Critical": 0, "Medium": 1, "Low": 2}
# endregion


# ==============================================================================
# region DOCX PARSER  (python-docx replacement for Aspose.Words)
# ==============================================================================
def _iter_block_items(parent):
    """Yield Paragraph and Table objects in document order (python-docx has no built-in)."""
    if isinstance(parent, _DocxDocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _is_heading_style(style_name: str) -> bool:
    if not style_name:
        return False
    s = style_name.lower()
    if s.startswith("toc"):
        return False
    if style_name.startswith("Heading") or "heading" in s:
        return True
    if style_name.startswith("Title") or style_name.startswith("Subtitle"):
        return True
    return False


def _clean_heading_text(text: str) -> str:
    text = re.sub(r"[\x00-\x1F]", "", text)
    text = re.sub(r"\s+", " ", text.strip())
    return text if text.strip() else "Untitled Section"


def _extract_table_text(table: Table) -> str:
    lines = ["[TABLE]"]
    for row in table.rows:
        cells = [cell.text.strip().replace("\a", "") for cell in row.cells]
        lines.append(" | ".join(cells))
    lines.append("[/TABLE]")
    return "\n".join(lines) + "\n"


def _process_image_for_llm(original_bytes: bytes, mime_type: str) -> tuple[bytes, int, int, bool]:
    """Resize image if larger than MAX_IMAGE_PIXELS while keeping aspect ratio."""
    max_pixels = CONFIG.MAX_IMAGE_PIXELS
    try:
        with Image.open(io.BytesIO(original_bytes)) as img:
            ow, oh = img.width, img.height
            if ow <= max_pixels and oh <= max_pixels:
                return original_bytes, ow, oh, False

            ratio = min(max_pixels / ow, max_pixels / oh)
            nw, nh = int(ow * ratio), int(oh * ratio)
            print(f"    \U0001f4f7 Resizing image from {ow}x{oh} to {nw}x{nh}")
            resized = img.resize((nw, nh), Image.LANCZOS)

            out = io.BytesIO()
            if "jpeg" in mime_type.lower() or "jpg" in mime_type.lower():
                if resized.mode in ("RGBA", "P"):
                    resized = resized.convert("RGB")
                resized.save(out, format="JPEG", quality=95)
            else:
                resized.save(out, format="PNG", compress_level=1)
            return out.getvalue(), ow, oh, True
    except Exception as ex:  # noqa: BLE001
        print(f"    \u26a0\ufe0f Image processing failed: {ex}")
        return original_bytes, 0, 0, False


def _extract_images_from_block(block, doc_part, current_section: ParsedSection, global_index: list[int]) -> None:
    """Find embedded images inside a paragraph/table element and append to the section."""
    element = block._p if isinstance(block, Paragraph) else block._tbl
    blips = element.findall(".//" + qn("a:blip"))
    for blip in blips:
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid or rid not in doc_part.related_parts:
            continue
        try:
            image_part = doc_part.related_parts[rid]
            blob = image_part.blob
            content_type = getattr(image_part, "content_type", "") or "image/png"
            if not content_type.startswith("image/"):
                content_type = "image/png"
            processed, ow, oh, was_resized = _process_image_for_llm(blob, content_type)
            current_section.images.append(
                ImageInfo(
                    data=processed,
                    mime_type=content_type,
                    page_no=0,
                    image_index=global_index[0],
                    alt_text="",
                    original_width=ow,
                    original_height=oh,
                    was_resized=was_resized,
                )
            )
            global_index[0] += 1
            print(f"    \U0001f4f7 Image {global_index[0]} extracted -> Section: '{current_section.section_name}'")
        except Exception as ex:  # noqa: BLE001
            print(f"    \u26a0\ufe0f Failed to extract image: {ex}")


def parse_docx(file_path: str) -> list[ParsedSection]:
    """Parse a Word document into sections based on heading styles."""
    doc = DocxDocument(file_path)
    doc_part = doc.part
    sections: list[ParsedSection] = []

    current_section: Optional[ParsedSection] = None
    para_buffer: list[str] = []
    paragraph_index = [0]
    global_image_index = [0]

    def flush_paragraph():
        if not para_buffer:
            return
        text = "\n".join(para_buffer).strip()
        para_buffer.clear()
        if text and current_section is not None:
            current_section.paragraphs.append(
                ParagraphInfo(page_no=0, para_index=paragraph_index[0], para_text=text, character_count=len(text))
            )
            paragraph_index[0] += 1

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            para_text = block.text.strip()

            if not para_text or para_text == "\f":
                if CONFIG.ENABLE_IMAGE_ANALYSIS and current_section is not None:
                    _extract_images_from_block(block, doc_part, current_section, global_image_index)
                continue

            try:
                style_name = block.style.name if block.style is not None else ""
            except Exception:  # noqa: BLE001
                style_name = ""

            if _is_heading_style(style_name):
                flush_paragraph()
                current_section = ParsedSection(section_name=_clean_heading_text(para_text))
                sections.append(current_section)
                paragraph_index[0] = 0
            else:
                if current_section is None:
                    current_section = ParsedSection(section_name="Document_Start")
                    sections.append(current_section)
                para_buffer.append(para_text)
                if sum(len(p) for p in para_buffer) > 2000:
                    flush_paragraph()

            if CONFIG.ENABLE_IMAGE_ANALYSIS and current_section is not None:
                _extract_images_from_block(block, doc_part, current_section, global_image_index)

        elif isinstance(block, Table):
            if current_section is None:
                current_section = ParsedSection(section_name="Document_Start")
                sections.append(current_section)
            table_text = _extract_table_text(block)
            if table_text.strip():
                para_buffer.append(table_text)
                flush_paragraph()
            if CONFIG.ENABLE_IMAGE_ANALYSIS:
                _extract_images_from_block(block, doc_part, current_section, global_image_index)

    flush_paragraph()
    return [s for s in sections if s.paragraphs or s.images]
# endregion


# ==============================================================================
# region MARKDOWN / TEXT PARSER  (fallback for .md / .txt)
# ==============================================================================
def _is_all_caps(text: str) -> bool:
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def _is_title_case(text: str) -> bool:
    words = [w for w in text.split(" ") if w]
    return bool(words) and words[0][0].isupper()


def parse_markdown(content: str) -> list[ParsedSection]:
    if not content or not content.strip():
        return []

    lines = content.split("\n")
    sections: list[ParsedSection] = []
    current_section = ParsedSection(section_name="User Selection")
    sections.append(current_section)

    para_buffer: list[str] = []
    paragraph_index = [0]
    found_heading = False

    def flush_paragraph():
        if not para_buffer:
            return
        text = "\n".join(para_buffer).strip()
        para_buffer.clear()
        if text:
            current_section.paragraphs.append(
                ParagraphInfo(page_no=0, para_index=paragraph_index[0], para_text=text, character_count=len(text))
            )
            paragraph_index[0] += 1

    for raw_line in lines:
        line = raw_line.rstrip()

        if re.match(r"^#{1,6}\s", line):
            flush_paragraph()
            found_heading = True
            heading_text = line.lstrip("#").strip()
            current_section = ParsedSection(section_name=heading_text if heading_text else "Untitled Section")
            sections.append(current_section)
            paragraph_index[0] = 0
            continue

        cleaned = re.sub(r"!\[[^\]]*\]\([^)]*\)|!\[[^\]]*\]", "", line.strip()).strip()
        bold_match = re.match(r"^\*\*([^*]+)\*\*:?$", cleaned)
        if bold_match:
            bold_text = bold_match.group(1).strip()
            bold_text = re.sub(r"\+\+([^+]+)\+\+", r"\1", bold_text).strip()
            if 3 <= len(bold_text) <= 100 and (_is_all_caps(bold_text) or _is_title_case(bold_text)):
                flush_paragraph()
                found_heading = True
                current_section = ParsedSection(section_name=bold_text)
                sections.append(current_section)
                paragraph_index[0] = 0
                continue

        if not line.strip():
            flush_paragraph()
            continue

        para_buffer.append(line)

    flush_paragraph()
    sections = [s for s in sections if s.paragraphs or s.images]

    if found_heading and sections and sections[0].section_name == "User Selection":
        sections[0].section_name = "Document Preamble"

    return sections
# endregion


# ==============================================================================
# region PDF PARSER  (PyMuPDF / fitz)
# ==============================================================================
def parse_pdf(file_path: str) -> list[ParsedSection]:
    """Parse a PDF document into sections.

    Uses font-size heuristics to detect headings: any line whose font size is
    larger than the document's median body size (or is bold/all-caps and short)
    is treated as a section heading — similar to how the markdown parser handles
    bold lines.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    sections: list[ParsedSection] = []
    current_section: Optional[ParsedSection] = None
    para_buffer: list[str] = []
    paragraph_index = [0]

    # First pass: collect font sizes to find the median body size.
    all_sizes: list[float] = []
    for page in doc:
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0:  # text block
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        all_sizes.append(span.get("size", 12))

    # Median body size — headings are typically >= 1.2x this
    if all_sizes:
        sorted_sizes = sorted(all_sizes)
        median_size = sorted_sizes[len(sorted_sizes) // 2]
    else:
        median_size = 12.0
    heading_threshold = median_size * 1.2

    def flush_paragraph():
        if not para_buffer:
            return
        text = "\n".join(para_buffer).strip()
        para_buffer.clear()
        if text and current_section is not None:
            current_section.paragraphs.append(
                ParagraphInfo(page_no=0, para_index=paragraph_index[0], para_text=text, character_count=len(text))
            )
            paragraph_index[0] += 1

    def is_heading(text: str, size: float, flags: int) -> bool:
        if not text.strip() or len(text.strip()) > 120:
            return False
        # Large font
        if size >= heading_threshold:
            return True
        # Bold (bit 4 in fitz flags) and short
        is_bold = bool(flags & (1 << 4))
        if is_bold and len(text.strip()) <= 80:
            return True
        # ALL CAPS short line
        if _is_all_caps(text.strip()) and len(text.strip()) <= 80:
            return True
        return False

    # Second pass: extract text with section splitting.
    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_text_parts: list[str] = []
                line_size = 12.0
                line_flags = 0
                for span in line.get("spans", []):
                    span_text = span.get("text", "")
                    if span_text:
                        line_text_parts.append(span_text)
                        line_size = max(line_size, span.get("size", 12.0))
                        line_flags |= span.get("flags", 0)

                line_text = "".join(line_text_parts).strip()
                if not line_text:
                    flush_paragraph()
                    continue

                if is_heading(line_text, line_size, line_flags):
                    flush_paragraph()
                    current_section = ParsedSection(section_name=_clean_heading_text(line_text))
                    sections.append(current_section)
                    paragraph_index[0] = 0
                else:
                    if current_section is None:
                        current_section = ParsedSection(section_name="Document_Start")
                        sections.append(current_section)
                    para_buffer.append(line_text)
                    if sum(len(p) for p in para_buffer) > 2000:
                        flush_paragraph()

        flush_paragraph()

    doc.close()
    return [s for s in sections if s.paragraphs or s.images]
# endregion


# ==============================================================================
# region SECTION CONVERSION (+ splitting)
# ==============================================================================
def _split_section_by_paragraphs(section: ParsedSection, max_chars: int) -> list[tuple[str, list[ParagraphInfo]]]:
    chunks: list[tuple[str, list[ParagraphInfo]]] = []
    cur_paras: list[ParagraphInfo] = []
    cur_len = 0
    cur_parts: list[str] = []

    for para in section.paragraphs:
        if cur_len + len(para.para_text) > max_chars and cur_paras:
            chunks.append(("\n".join(cur_parts).strip(), list(cur_paras)))
            cur_paras, cur_parts, cur_len = [], [], 0
        cur_paras.append(para)
        cur_parts.append(para.para_text)
        cur_len += len(para.para_text)

    if cur_paras:
        chunks.append(("\n".join(cur_parts).strip(), cur_paras))
    return chunks


def convert_to_processable_sections(parsed_sections: list[ParsedSection]) -> list[ProcessableSection]:
    result: list[ProcessableSection] = []
    section_id = 1

    for section in parsed_sections:
        full_content = "\n\n".join(p.para_text for p in section.paragraphs)
        char_count = len(full_content)

        pages = [p.page_no for p in section.paragraphs if p.page_no > 0]
        start_page = min(pages) if pages else 0
        end_page = max(pages) if pages else 0

        if char_count > CONFIG.MAX_SECTION_CHARACTERS:
            chunks = _split_section_by_paragraphs(section, CONFIG.MAX_SECTION_CHARACTERS)
            for chunk_index, (content, paras) in enumerate(chunks):
                cpages = [p.page_no for p in paras if p.page_no > 0]
                result.append(
                    ProcessableSection(
                        id=f"{section_id}-{chunk_index}",
                        title=section.section_name,
                        content=content,
                        character_count=len(content),
                        is_split=True,
                        split_index=chunk_index,
                        total_splits=len(chunks),
                        paragraphs=paras,
                        start_page=min(cpages) if cpages else 0,
                        end_page=max(cpages) if cpages else 0,
                        images=list(section.images) if chunk_index == 0 else [],
                    )
                )
        else:
            result.append(
                ProcessableSection(
                    id=str(section_id),
                    title=section.section_name,
                    content=full_content,
                    character_count=char_count,
                    is_split=False,
                    split_index=0,
                    total_splits=1,
                    paragraphs=list(section.paragraphs),
                    start_page=start_page,
                    end_page=end_page,
                    images=list(section.images),
                )
            )
        section_id += 1

    return result


def generate_user_friendly_message(inp: DocumentJob, sections: list[ProcessableSection]) -> str:
    count = len(sections)
    names = list(dict.fromkeys(s.title for s in sections))
    if inp.is_full_document:
        extra = f" and {len(names) - 5} more..." if len(names) > 5 else ""
        return f"Processing {count} sections from your document: {', '.join(names[:5])}{extra}"
    if inp.user_selection:
        return f"Analyzing your highlighted selection ({count} section(s) identified)..."
    if inp.selected_sections:
        return f"Reviewing your selected sections: {', '.join(inp.selected_sections)}"
    return f"Processing {count} section(s)..."
# endregion


# ==============================================================================
# region EXECUTOR 1 — DocumentSegmenter
# ==============================================================================
class DocumentSegmenterExecutor(Executor):
    def __init__(self):
        super().__init__(id="DocumentSegmenter")

    @handler
    async def run(self, inp: DocumentJob, ctx: WorkflowContext[DispatchSignal]) -> None:
        print("[1] DocumentSegmenterExecutor - Parsing document")
        if not inp.is_valid:
            raise ValueError("Either file_path or selection_text must be provided.")

        if inp.selection_text and inp.selection_text.strip():
            print("    Mode: User Selection (direct text)")
            parsed = parse_markdown(inp.selection_text)
            print(f"    \u2713 Parsed {len(parsed)} sections from selection text")
            extracted = convert_to_processable_sections(parsed)
        else:
            file_path = inp.file_path or ""
            ext = Path(file_path).suffix.lower()
            print(f"    File: {file_path}")
            if ext in (".docx", ".doc"):
                print("    Using: python-docx (heading detection)")
                parsed = parse_docx(file_path)
                print(f"    \u2713 Parsed {len(parsed)} sections from Word document")
                extracted = convert_to_processable_sections(parsed)
            elif ext == ".pdf":
                print("    Using: PyMuPDF (PDF text extraction)")
                parsed = parse_pdf(file_path)
                print(f"    \u2713 Parsed {len(parsed)} sections from PDF")
                extracted = convert_to_processable_sections(parsed)
            elif ext in (".md", ".markdown", ".txt"):
                print("    Using: Markdown/Text parser")
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                parsed = parse_markdown(content)
                print(f"    \u2713 Parsed {len(parsed)} sections from Markdown")
                extracted = convert_to_processable_sections(parsed)
            else:
                raise ValueError(f"Unsupported file format: {ext}. Use .docx, .pdf, .md, or .txt files.")

        print(f"    \u2713 Created {len(extracted)} processable sections")
        for s in extracted:
            split_info = f" [Split {s.split_index + 1}/{s.total_splits}]" if s.is_split else ""
            print(f"       - {s.title}{split_info}: {s.character_count:,} chars, {len(s.paragraphs)} paragraphs, {len(s.images)} images")

        message = generate_user_friendly_message(inp, extracted)
        print(f"    \U0001f4dd {message}")

        await ctx.send_message(
            DispatchSignal(
                initial_sections=ExtractedSections(
                    sections=extracted,
                    total_count=len(extracted),
                    user_friendly_message=message,
                ),
                resume_from_section=inp.resume_from_section,
            )
        )
# endregion


# ==============================================================================
# region EXECUTOR 2 — SegmentDispatcher (loop driver)
class SectionProgressEvent(WorkflowEvent):
    """Live progress signal emitted as each section starts being analyzed.

    `data` carries: {"index", "total", "title", "status"}.
    Surfaced to the API via workflow.run_stream() so the UI can show which
    section is currently being processed.
    """


# ==============================================================================
class SegmentDispatcherExecutor(Executor):
    def __init__(self):
        super().__init__(id="SegmentDispatcher")
        self._current_index = 0
        self._sections: list[ProcessableSection] = []
        self._results: list[ValidationResult] = []
        self._initialized = False

    @handler
    async def run(self, inp: DispatchSignal, ctx: WorkflowContext[DispatchResult]) -> None:
        # Initialise on first call
        if inp.initial_sections is not None and not self._initialized:
            self._sections = inp.initial_sections.sections
            self._initialized = True
            if 0 < inp.resume_from_section < len(self._sections):
                self._current_index = inp.resume_from_section
                print(f"[2] SegmentDispatcherExecutor - Resuming from section {self._current_index + 1}/{len(self._sections)}")
            else:
                self._current_index = 0

        # Collect the validation result from the loop-back signal
        if inp.loop_back_signal is not None:
            self._results.append(inp.loop_back_signal)

        if self._current_index >= len(self._sections):
            print(f"[2] SegmentDispatcherExecutor - All {len(self._sections)} sections processed")
            print(f"    Retrieved {len(self._results)} validation results for aggregation")
            await ctx.send_message(
                DispatchResult(
                    current_section=None,
                    is_complete=True,
                    processed_count=len(self._sections),
                    total_count=len(self._sections),
                    all_validation_results=list(self._results),
                )
            )
            return

        current = self._sections[self._current_index]
        print(f"[2] SegmentDispatcherExecutor - Emitting section {self._current_index + 1}/{len(self._sections)}: {current.title}")
        self._current_index += 1

        # Live progress signal for the UI (which section is being analyzed now).
        await ctx.add_event(
            SectionProgressEvent(
                {
                    "index": self._current_index,
                    "total": len(self._sections),
                    "title": current.title,
                    "status": "analyzing",
                }
            )
        )

        await ctx.send_message(
            DispatchResult(
                current_section=current,
                is_complete=False,
                processed_count=self._current_index,
                total_count=len(self._sections),
                all_validation_results=[],
            )
        )
# endregion


# ==============================================================================
# region EXECUTOR 3 — ReferenceRetriever (Verba / Weaviate vector retrieval)
# ==============================================================================
class ReferenceRetrieverExecutor(Executor):
    def __init__(self):
        super().__init__(id="ReferenceRetriever")
        self._client: Optional[VerbaClient] = None

    def _get_client(self) -> VerbaClient:
        if self._client is None:
            self._client = VerbaClient.from_env()
        return self._client

    @handler
    async def run(self, inp: DispatchResult, ctx: WorkflowContext[RetrievalResult]) -> None:
        section = inp.current_section
        assert section is not None
        print(f"[3] ReferenceRetrieverExecutor - Querying Verba for section: {section.title}")
        print(f"    Section size: {section.character_count:,} characters")

        chunks = self._build_search_chunks(section)
        print(f"    Split into {len(chunks)} search chunk(s)")

        all_messages: list[ReferenceSnippet] = []
        for i, chunk in enumerate(chunks, start=1):
            print(f"    Chunk {i}: {len(chunk):,} characters")
            all_messages.extend(await self._search(chunk))

        print(f"    Total retrieved before dedup: {len(all_messages)} reference chunks")
        deduped = self._deduplicate(all_messages)
        print(f"    \u2713 After deduplication: {len(deduped)} unique reference chunks")
        for km in deduped[:3]:
            snippet = (km.content[:100] + "...") if len(km.content) > 100 else km.content
            print(f"       [{km.id}] {km.source} - {snippet}")

        await ctx.send_message(RetrievalResult(section=section, retrieved_references=deduped))

    async def _search(self, query: str) -> list[ReferenceSnippet]:
        """Send a section query to Verba and map the retrieved chunks to ReferenceSnippet."""
        try:
            results = await self._get_client().query(query, top_k=CONFIG.VERBA_TOP_K)
        except Exception as ex:  # noqa: BLE001
            print(f"    \u274c Verba query error: {ex}")
            return []

        messages: list[ReferenceSnippet] = []
        for idx, result in enumerate(results):
            content = (result.get("content") or "").strip()
            if not content:
                continue
            messages.append(
                ReferenceSnippet(id=f"km-{idx:03d}", content=content, source=result.get("source", ""))
            )
        return messages

    @staticmethod
    def _build_search_chunks(section: ProcessableSection) -> list[str]:
        chunks: list[str] = []
        title_prefix = f"{section.title}\n\n"
        available = CONFIG.MAX_SEARCH_QUERY_CHARACTERS - len(title_prefix) - 100

        cur_parts = [title_prefix]
        cur_len = 0

        for para in section.paragraphs:
            if cur_len + len(para.para_text) <= available:
                cur_parts.append(para.para_text + "\n")
                cur_len += len(para.para_text)
            else:
                if cur_len > 0:
                    chunks.append("".join(cur_parts).strip())
                cur_parts = [title_prefix]
                cur_len = 0
                if len(para.para_text) <= available:
                    cur_parts.append(para.para_text + "\n")
                    cur_len = len(para.para_text)
                else:
                    offset = 0
                    while offset < len(para.para_text):
                        seg = para.para_text[offset: offset + available]
                        chunks.append(f"{title_prefix}{seg}".strip())
                        offset += available
                    cur_parts = [title_prefix]
                    cur_len = 0

        if cur_len > 0:
            chunks.append("".join(cur_parts).strip())
        if not chunks:
            chunks.append(section.title)
        return chunks

    @staticmethod
    def _deduplicate(messages: list[ReferenceSnippet]) -> list[ReferenceSnippet]:
        seen: set[str] = set()
        out: list[ReferenceSnippet] = []
        for msg in messages:
            normalized = re.sub(r"\s+", " ", msg.content.strip()).lower()
            if not normalized:
                continue
            key = normalized[:200]
            if key not in seen:
                seen.add(key)
                out.append(msg)
        return out
# endregion


# ==============================================================================
# region EXECUTOR 4 — AlignmentReviewer (Azure OpenAI, multimodal)
# ==============================================================================
class AlignmentReviewerExecutor(Executor):
    def __init__(self):
        super().__init__(id="AlignmentReviewer")
        self._chat_client: Optional[AzureOpenAIChatClient] = None
        self._agent = None

    def _ensure_agent(self):
        if self._chat_client is None:
            self._chat_client = AzureOpenAIChatClient(
                endpoint=CONFIG.AZURE_OPENAI_ENDPOINT,
                api_key=CONFIG.AZURE_OPENAI_API_KEY,
                deployment_name=CONFIG.AZURE_OPENAI_DEPLOYMENT,
                api_version=CONFIG.AZURE_OPENAI_API_VERSION,
            )
            self._agent = self._chat_client.as_agent(
                name=CONFIG.REVIEW_AGENT_NAME,
                instructions=CONFIG.REVIEW_AGENT_INSTRUCTIONS,
            )

    @handler
    async def run(self, inp: RetrievalResult, ctx: WorkflowContext[DispatchSignal]) -> None:
        section = inp.section
        print(f"[4] AlignmentReviewerExecutor - Validating section: {section.title}")
        print(f"    Section: {section.character_count:,} chars | References: {len(inp.retrieved_references)} | Images: {len(section.images)}")

        self._ensure_agent()
        prompt = self._build_prompt(section, inp.retrieved_references)
        print(f"    Prompt size: {len(prompt):,} characters")
        if section.images:
            print(f"    \U0001f4f7 Including {len(section.images)} image(s) in multimodal analysis")

        try:
            result = await self._call_llm(section, prompt)
            print(f"    \u2713 LLM Response - Aligned: {result.is_aligned}, Confidence: {result.confidence:.0%}, Severity: {result.severity}")
            if result.discrepancies:
                print(f"    \u26a0\ufe0f  Found {len(result.discrepancies)} discrepancies")
        except Exception as ex:  # noqa: BLE001
            print(f"    \u274c LLM Error: {ex}")
            result = ValidationResult(
                section_id=section.id,
                section_title=section.title,
                is_aligned=True,
                confidence=0.5,
                severity="Medium",
                discrepancies=[],
                reasoning=f"Validation error: {ex}. Manual review recommended.",
            )

        await ctx.send_message(DispatchSignal(loop_back_signal=result))

    @staticmethod
    def _build_prompt(section: ProcessableSection, references: list[ReferenceSnippet]) -> str:
        b: list[str] = []
        b.append("## SECTION TO VALIDATE")
        b.append(f"**Section Title:** {section.title}")
        b.append(f"**Pages:** {section.start_page}-{section.end_page}")
        b.append("")
        b.append("**Content:**")
        b.append(section.content)
        b.append("")

        if section.images:
            b.append("**Images in this section:**")
            for i, img in enumerate(section.images):
                alt = f" (Alt: {img.alt_text})" if img.alt_text.strip() else ""
                resize = f" [resized from {img.original_width}x{img.original_height}]" if img.was_resized else ""
                b.append(f"- Image {i + 1}: Page {img.page_no}{alt}{resize}")
            b.append("")
            b.append("IMPORTANT FOR IMAGE ANALYSIS:")
            b.append("- Analyze the images for any visual content (text, charts, diagrams, labels) that may contradict the reference content")
            b.append("- Report image discrepancies using the SAME discrepancy format as text discrepancies")
            b.append("- For image discrepancies, set 'isFromImage': true and include:")
            b.append("  - 'imageDescription': describe what you see in the image that contradicts the reference content")
            b.append("  - 'imageReference': which image (e.g., 'Image 1 on Page 3')")
            b.append("  - 'statementInDocument': quote the exact text visible in the image that contradicts")
            b.append("  - 'conflictingReference': the reference content it contradicts")
            b.append("")

        b.append("---")
        b.append("")
        b.append("## APPROVED REFERENCE CONTENT (Reference)")
        b.append("Compare the section content against this approved reference content:")
        b.append("")
        for i, km in enumerate(references):
            b.append(f"### Reference {i + 1}")
            b.append(km.content)
            b.append("")
        b.append("---")
        b.append("")
        b.append("Analyze the section content (including any images provided) against the reference content and output your assessment as JSON.")
        return "\n".join(b)

    async def _call_llm(self, section: ProcessableSection, prompt: str) -> ValidationResult:
        if section.images:
            contents: list[Content] = [Content.from_text(prompt)]
            for image in section.images:
                try:
                    contents.append(Content.from_data(data=image.data, media_type=image.mime_type))
                except Exception as ex:  # noqa: BLE001
                    print(f"    \u26a0\ufe0f Failed to add image to message: {ex}")
            messages = [
                ChatMessage(role="system", text=CONFIG.REVIEW_AGENT_INSTRUCTIONS),
                ChatMessage(role="user", contents=contents),
            ]
            response = await self._chat_client.get_response(messages)
            json_response = response.text or "{}"
        else:
            response = await self._agent.run(prompt)
            json_response = response.text or "{}"

        json_response = self._clean_json(json_response)
        return self._parse_response(json_response, section)

    @staticmethod
    def _clean_json(response: str) -> str:
        response = response.strip()
        if response.lower().startswith("```json"):
            response = response[7:]
        elif response.startswith("```"):
            response = response[3:]
        if response.endswith("```"):
            response = response[:-3]
        return response.strip()

    @staticmethod
    def _parse_response(json_response: str, section: ProcessableSection) -> ValidationResult:
        try:
            data = json.loads(json_response)

            severity_raw = (data.get("severity") or "").lower()
            severity = {
                "critical": "Critical",
                "high": "Critical",
                "medium": "Medium",
                "low": "Low",
            }.get(severity_raw, "Medium")

            raw_discrepancies = data.get("discrepancies") or []
            validated: list[Discrepancy] = []

            for d in raw_discrepancies:
                statement = (d.get("statementInDocument") or "").strip()
                conflicting = (d.get("conflictingReference") or "").strip()
                if not statement or not conflicting:
                    print("    \u26a0\ufe0f Skipping discrepancy: empty statementInDocument or conflictingReference")
                    continue

                if d.get("isFromImage"):
                    print(f"    \U0001f4f7 Image discrepancy found: {d.get('imageReference') or 'Image'}")
                    validated.append(
                        Discrepancy(
                            statement_in_document=statement,
                            conflicting_reference=conflicting,
                            issue_description=d.get("issueDescription") or "",
                            is_from_image=True,
                            image_description=d.get("imageDescription") or "",
                            image_reference=d.get("imageReference") or "",
                        )
                    )
                    continue

                found = statement.lower() in section.content.lower()
                if not found:
                    words = [w for w in statement.split(" ") if len(w) > 3]
                    matched = sum(1 for w in words if w.lower() in section.content.lower())
                    ratio = (matched / len(words)) if words else 0.0
                    if ratio < 0.5:
                        print("    \u26a0\ufe0f INVALID DISCREPANCY - statementInDocument NOT FOUND in section (skipping)")
                        continue
                    print(f"    \u2139\ufe0f Partial match ({ratio:.0%}) - keeping discrepancy with warning")

                validated.append(
                    Discrepancy(
                        statement_in_document=statement,
                        conflicting_reference=conflicting,
                        issue_description=d.get("issueDescription") or "",
                        is_from_image=False,
                        image_description="",
                        image_reference="",
                    )
                )

            has_valid = len(validated) > 0
            final_aligned = bool(data.get("isAligned")) if has_valid else True

            return ValidationResult(
                section_id=section.id,
                section_title=section.title,
                is_aligned=final_aligned,
                confidence=float(data.get("confidence", 0.0)),
                severity=severity,
                discrepancies=validated,
                reasoning=data.get("reasoning") or "",
            )
        except (ValueError, TypeError) as ex:
            print(f"    \u26a0\ufe0f JSON Parse Error: {ex}")
            print(f"    Raw response: {json_response[:200]}...")
            return ValidationResult(
                section_id=section.id,
                section_title=section.title,
                is_aligned=True,
                confidence=0.6,
                severity="Medium",
                discrepancies=[],
                reasoning=f"Could not parse LLM response. Manual review recommended. Error: {ex}",
            )
# endregion


# ==============================================================================
# region EXECUTOR 5 — ReportCompiler (final QC report)
# ==============================================================================
class ReportCompilerExecutor(Executor):
    def __init__(self):
        super().__init__(id="ReportCompiler")

    @handler
    async def run(self, inp: DispatchResult, ctx: WorkflowContext[Never, dict]) -> None:
        print(f"[5] ReportCompilerExecutor - Aggregating results (was_paused: {inp.was_paused})")

        if inp.was_paused:
            print(f"   Workflow was paused at section {inp.processed_count}/{inp.total_count}")
            await ctx.yield_output(
                {
                    "totalSections": inp.total_count,
                    "cleanSections": 0,
                    "issuesFound": 0,
                    "discrepancies": [],
                    "processingTime": "",
                    "wasPaused": True,
                    "completedSections": inp.processed_count,
                }
            )
            return

        all_results = inp.all_validation_results or []
        print(f"   Processing {len(all_results)} validation results")

        clean = [r for r in all_results if r.is_aligned]
        discrepancy_sections = [r for r in all_results if not r.is_aligned]

        # Mirrors C#: OrderByDescending(Severity).ThenByDescending(Confidence)
        discrepancy_sections.sort(
            key=lambda r: (SEVERITY_RANK.get(r.severity, 1), r.confidence),
            reverse=True,
        )

        discrepancy_results = [
            {
                "sectionId": r.section_id,
                "sectionTitle": r.section_title,
                "confidence": r.confidence,
                "severity": r.severity,
                "discrepancies": [
                    {
                        "statementInDocument": d.statement_in_document,
                        "conflictingReference": d.conflicting_reference,
                        "issueDescription": d.issue_description,
                        "isFromImage": d.is_from_image,
                        "imageDescription": d.image_description,
                        "imageReference": d.image_reference,
                    }
                    for d in r.discrepancies
                ],
                "reasoning": r.reasoning,
            }
            for r in discrepancy_sections
        ]

        report = {
            "totalSections": len(all_results),
            "cleanSections": len(clean),
            "issuesFound": len(discrepancy_sections),
            "discrepancies": discrepancy_results,
            "processingTime": "",
            "wasPaused": False,
            "completedSections": 0,
        }

        print("   \u2713 Report generated:")
        print(f"      Total sections: {report['totalSections']}")
        print(f"      Clean: {report['cleanSections']}")
        print(f"      Issues: {report['issuesFound']}")

        await ctx.yield_output(report)
# endregion


# ==============================================================================
# region WORKFLOW BUILDER
# ==============================================================================
def build_workflow():
    segmenter = DocumentSegmenterExecutor()
    dispatcher = SegmentDispatcherExecutor()
    retriever = ReferenceRetrieverExecutor()
    reviewer = AlignmentReviewerExecutor()
    compiler = ReportCompilerExecutor()

    builder = WorkflowBuilder(max_iterations=100000)
    builder.set_start_executor(segmenter)

    # Segmenter -> Dispatcher
    builder.add_edge(segmenter, dispatcher)
    # Dispatcher -> Retriever (only while sections remain)
    builder.add_edge(dispatcher, retriever, condition=lambda r: isinstance(r, DispatchResult) and not r.is_complete)
    # Retriever -> Reviewer
    builder.add_edge(retriever, reviewer)
    # Reviewer -> Dispatcher (loop back)
    builder.add_edge(reviewer, dispatcher)
    # Dispatcher -> Compiler (when complete)
    builder.add_edge(dispatcher, compiler, condition=lambda r: isinstance(r, DispatchResult) and r.is_complete)

    return builder.build()
# endregion


# ==============================================================================
# region OUTPUT + ENTRY POINT
# ==============================================================================
def resolve_document_path() -> str:
    """Find a document in INPUT (preferred), else fall back to the markdown/text file."""
    CONFIG.INPUT_FOLDER.mkdir(parents=True, exist_ok=True)

    # Priority: .docx > .pdf > .md/.txt
    docx_candidates = sorted(
        p for p in CONFIG.INPUT_FOLDER.glob("*.docx") if not p.name.startswith("~$")
    )
    if docx_candidates:
        print(f"\U0001f4c4 Found Word document: {docx_candidates[0]}")
        return str(docx_candidates[0])

    pdf_candidates = sorted(CONFIG.INPUT_FOLDER.glob("*.pdf"))
    if pdf_candidates:
        print(f"\U0001f4c4 Found PDF document: {pdf_candidates[0]}")
        return str(pdf_candidates[0])

    if CONFIG.MARKDOWN_FALLBACK_PATH.exists():
        print(f"\U0001f4c4 Found Markdown file: {CONFIG.MARKDOWN_FALLBACK_PATH} (fallback mode)")
        return str(CONFIG.MARKDOWN_FALLBACK_PATH)

    raise FileNotFoundError(
        f"No document found. Place a .docx or .pdf file in: {CONFIG.INPUT_FOLDER} "
        f"(or provide {CONFIG.MARKDOWN_FALLBACK_PATH})"
    )


def write_report(report: dict, elapsed_seconds: float) -> Path:
    CONFIG.OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)
    report = dict(report)
    report["processingTime"] = (
        f"{elapsed_seconds:.1f} seconds" if elapsed_seconds < 60 else f"{elapsed_seconds / 60:.1f} minutes"
    )

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_path = CONFIG.OUTPUT_FOLDER / f"QC_Report_{timestamp}.json"
    content = json.dumps(report, indent=2, ensure_ascii=False)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"\n\U0001f4be QC Report saved to: {file_path}")
    print(f"   File size: {len(content):,} characters")
    return file_path


def export_diagram(workflow) -> None:
    """Best-effort Mermaid/DOT diagram export (non-fatal)."""
    try:
        from agent_framework import WorkflowViz

        CONFIG.DIAGRAM_OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)
        viz = WorkflowViz(workflow)
        mermaid = viz.to_mermaid()
        (CONFIG.DIAGRAM_OUTPUT_FOLDER / "workflow-diagram.mmd").write_text(mermaid, encoding="utf-8")
        print(f"\u2705 Mermaid diagram saved to {CONFIG.DIAGRAM_OUTPUT_FOLDER / 'workflow-diagram.mmd'}")
    except Exception as ex:  # noqa: BLE001
        print(f"\u2139\ufe0f Diagram export skipped: {ex}")


async def run_validation_for_file(document_path: str, write_to_disk: bool = True) -> dict:
    """Run the full document governance workflow for a single document and return the QC report.

    This is the programmatic entry point used by the API. It bypasses the
    input-folder auto-discovery and validates exactly the file provided.
    """
    workflow = build_workflow()

    inp = DocumentJob(
        file_path=document_path,
        is_full_document=True,
        user_selection=False,
        selected_sections=None,
    )

    start = datetime.now()
    events = await workflow.run(inp)
    elapsed = (datetime.now() - start).total_seconds()

    outputs = events.get_outputs()
    if not outputs:
        raise RuntimeError("Workflow produced no output.")

    report = dict(outputs[0])
    report["processingTime"] = (
        f"{elapsed:.1f} seconds" if elapsed < 60 else f"{elapsed / 60:.1f} minutes"
    )

    if write_to_disk:
        try:
            write_report(report, elapsed)
        except Exception as ex:  # noqa: BLE001
            print(f"\u2139\ufe0f Could not persist report to disk: {ex}")

    return report


async def run_validation_stream(document_path: str):
    """Streaming variant of `run_validation_for_file`.

    Async generator that yields progress dicts as sections are analyzed, then a
    final {"type": "report", "report": {...}} message. Used by the API's SSE
    endpoint so the UI can show live per-section progress.
    """
    workflow = build_workflow()

    inp = DocumentJob(
        file_path=document_path,
        is_full_document=True,
        user_selection=False,
        selected_sections=None,
    )

    start = datetime.now()
    report: dict | None = None

    async for event in workflow.run_stream(inp):
        if isinstance(event, SectionProgressEvent):
            yield {"type": "progress", **event.data}
        elif isinstance(event, WorkflowOutputEvent):
            report = event.data

    if report is None:
        raise RuntimeError("Workflow produced no output.")

    elapsed = (datetime.now() - start).total_seconds()
    report = dict(report)
    report["processingTime"] = (
        f"{elapsed:.1f} seconds" if elapsed < 60 else f"{elapsed / 60:.1f} minutes"
    )

    yield {"type": "report", "report": report}


async def main() -> int:
    print("=== Document Governance Workflow (Python) ===\n")

    workflow = build_workflow()
    export_diagram(workflow)

    document_path = resolve_document_path()
    inp = DocumentJob(
        file_path=document_path,
        is_full_document=True,
        user_selection=False,
        selected_sections=None,
    )

    print("\n\U0001f9ea Running workflow...\n")
    start = datetime.now()

    try:
        events = await workflow.run(inp)
    except Exception as ex:  # noqa: BLE001
        print(f"\u26a0\ufe0f  Workflow execution failed: {ex}")
        import traceback

        traceback.print_exc()
        return 1

    elapsed = (datetime.now() - start).total_seconds()
    outputs = events.get_outputs()

    if not outputs:
        print("\u26a0\ufe0f  Workflow produced no output.")
        return 1

    report = outputs[0]
    print("\n\U0001f4e4 Workflow output received")
    print(f"   Total sections: {report.get('totalSections')}")
    print(f"   Clean: {report.get('cleanSections')}")
    print(f"   Issues: {report.get('issuesFound')}")

    write_report(report, elapsed)
    print("\n\u2705 Workflow completed successfully!")
    return 0


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))
# endregion
