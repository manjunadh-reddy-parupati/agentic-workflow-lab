"""
Local Weaviate RAG client.

Replaces the previous Azure AI Search retrieval with a local, open-source RAG
backend powered directly by Weaviate (the same vector DB that backs Verba).

Why talk to Weaviate directly instead of Verba's HTTP API?
  Verba is built as a self-contained UI app: its `/api/*` routes are guarded by
  a same-origin middleware (returns 403 to non-browser callers), importing is
  done over a WebSocket with a full RAG-config payload, and it defaults to an
  Ollama embedder. For a clean, headless integration we instead use Weaviate's
  stable Python client. Weaviate vectorizes text locally via the bundled
  `text2vec-transformers` module (see docker-compose.yml) — no Ollama or any
  external subscription required. The Verba UI on :8000 still works as an
  optional inspector over the same Weaviate instance.

Responsibilities (interface preserved for the workflow/backend):
  * import_document(path)  -> extract text, chunk it, (re)create the collection,
                              and store chunks in Weaviate (auto-vectorized).
  * query(text, top_k)     -> nearText search for the most relevant chunks.

Configuration (env vars, see `from_env`):
    WEAVIATE_HOST       (default: localhost)
    WEAVIATE_HTTP_PORT  (default: 8080)
    WEAVIATE_GRPC_PORT  (default: 50051)
    RAG_COLLECTION      (default: SourceOfTruth)
"""

import asyncio
import os
import re
from pathlib import Path
from typing import Any

import weaviate
from weaviate.classes.config import Configure, DataType, Property
from weaviate.classes.query import MetadataQuery

# Files we can extract text from for ingestion.
_TEXT_EXTENSIONS = {".md", ".markdown", ".txt"}
_PDF_EXTENSIONS = {".pdf"}

# Rough chunking targets (characters). Paragraph-aware, with light overlap.
_CHUNK_CHARS = 900
_CHUNK_OVERLAP = 150


class VerbaClient:
    """Headless RAG client backed directly by a local Weaviate instance."""

    def __init__(
        self,
        host: str = "localhost",
        http_port: int = 8080,
        grpc_port: int = 50051,
        collection: str = "SourceOfTruth",
    ) -> None:
        self.host = host
        self.http_port = http_port
        self.grpc_port = grpc_port
        self.collection = collection

    @classmethod
    def from_env(cls) -> "VerbaClient":
        return cls(
            host=os.environ.get("WEAVIATE_HOST", "localhost"),
            http_port=int(os.environ.get("WEAVIATE_HTTP_PORT", "8080")),
            grpc_port=int(os.environ.get("WEAVIATE_GRPC_PORT", "50051")),
            collection=os.environ.get("RAG_COLLECTION", "SourceOfTruth"),
        )

    # ----------------------------------------------------------- connection --
    def _connect(self) -> "weaviate.WeaviateClient":
        return weaviate.connect_to_local(
            host=self.host,
            port=self.http_port,
            grpc_port=self.grpc_port,
        )

    # --------------------------------------------------------------- ingest --
    async def import_document(self, file_path: str) -> dict[str, Any]:
        """Extract, chunk and store the Source-of-Truth file in Weaviate."""
        return await asyncio.to_thread(self._sync_import, file_path)

    def _sync_import(self, file_path: str) -> dict[str, Any]:
        path = Path(file_path)
        text = _extract_text(path)
        chunks = _chunk_text(text)
        if not chunks:
            raise ValueError(f"No extractable text found in '{path.name}'.")

        client = self._connect()
        try:
            # Fresh index per uploaded Source of Truth.
            if client.collections.exists(self.collection):
                client.collections.delete(self.collection)

            client.collections.create(
                name=self.collection,
                vector_config=Configure.Vectors.text2vec_transformers(),
                properties=[
                    Property(name="content", data_type=DataType.TEXT),
                    Property(name="source", data_type=DataType.TEXT),
                ],
            )

            coll = client.collections.get(self.collection)
            with coll.batch.dynamic() as batch:
                for chunk in chunks:
                    batch.add_object(
                        properties={"content": chunk, "source": path.name}
                    )

            failed = len(coll.batch.failed_objects)
        finally:
            client.close()

        return {
            "status": "ingested",
            "collection": self.collection,
            "chunks": len(chunks),
            "failed": failed,
        }

    # ---------------------------------------------------------------- query --
    async def query(self, text: str, top_k: int = 3) -> list[dict[str, Any]]:
        """Retrieve the most relevant Source-of-Truth chunks for a section.

        Returns a list of {"content": str, "source": str} dicts.
        """
        return await asyncio.to_thread(self._sync_query, text, top_k)

    def _sync_query(self, text: str, top_k: int) -> list[dict[str, Any]]:
        client = self._connect()
        try:
            if not client.collections.exists(self.collection):
                # Source of Truth not ingested yet -> no context available.
                return []

            coll = client.collections.get(self.collection)
            res = coll.query.near_text(
                query=text,
                limit=top_k,
                return_metadata=MetadataQuery(distance=True),
            )
        finally:
            client.close()

        out: list[dict[str, Any]] = []
        for obj in res.objects:
            content = obj.properties.get("content", "")
            if not isinstance(content, str) or not content.strip():
                continue
            source = obj.properties.get("source") or "SourceOfTruth"
            distance = getattr(obj.metadata, "distance", None)
            label = f"{source} (distance: {distance:.4f})" if distance is not None else str(source)
            out.append({"content": content, "source": label})

        return out


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def _extract_text(path: Path) -> str:
    """Extract plain text from a .docx / .pdf / .md / .txt Source-of-Truth file."""
    ext = path.suffix.lower()
    if ext in {".docx", ".doc"}:
        from docx import Document  # local import; python-docx is a dependency

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if ext in _PDF_EXTENSIONS:
        import fitz  # PyMuPDF

        text_parts: list[str] = []
        with fitz.open(str(path)) as pdf_doc:
            for page in pdf_doc:
                page_text = page.get_text("text")
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
        return "\n\n".join(text_parts)

    if ext in _TEXT_EXTENSIONS:
        return path.read_text(encoding="utf-8", errors="ignore")

    raise ValueError(f"Unsupported Source-of-Truth file type: '{ext}'.")


def _chunk_text(text: str) -> list[str]:
    """Split text into paragraph-aware, lightly overlapping character chunks."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    buffer = ""

    for para in paragraphs:
        if len(buffer) + len(para) + 1 <= _CHUNK_CHARS:
            buffer = f"{buffer}\n{para}" if buffer else para
            continue

        if buffer:
            chunks.append(buffer)
        # Carry a little overlap from the previous buffer for context continuity.
        tail = buffer[-_CHUNK_OVERLAP:] if buffer else ""

        if len(para) <= _CHUNK_CHARS:
            buffer = f"{tail}\n{para}".strip() if tail else para
        else:
            # Hard-split very long paragraphs.
            start = 0
            while start < len(para):
                piece = para[start : start + _CHUNK_CHARS]
                chunks.append(piece.strip())
                start += _CHUNK_CHARS - _CHUNK_OVERLAP
            buffer = ""

    if buffer.strip():
        chunks.append(buffer.strip())

    return [c for c in chunks if c.strip()]
